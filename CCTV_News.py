#!/usr/bin/env python
# coding: utf-8

# In[4]:


import requests
from bs4 import BeautifulSoup
import time
import datetime
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt,RGBColor


# In[54]:


today = time.strftime("%Y%m%d", time.localtime()) 

time_now = time.strftime('%H%M',time.localtime())

yest=datetime.datetime.now()+datetime.timedelta(days=-1)
yesterday = datetime.datetime.strftime(yest,"%Y%m%d")


# In[ ]:


if int(time_now) < 1930:
    today = yesterday


# In[3]:


xwlb_url = 'https://tv.cctv.com/lm/xwlb/day/{}.shtml'.format(today)
jsxsk = 'http://v.jstv.com/jsxsk'

xwlb_headers = {
"Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,image/apng,*/*;q=0.8,application/signed-exchange;v=b3;q=0.9",
"Accept-Encoding": "gzip, deflate, br",
"Accept-Language": "zh-CN,zh;q=0.9,en;q=0.8,en-GB;q=0.7,en-US;q=0.6",
"Cache-Control": "max-age=0",
"Connection": "keep-alive",
"Cookie": "float_title=%u6B63%u5728%u64AD%u51FA%uFF1A%u300A%u5929%u4E0B%u8DB3%u7403%u300B%u76D8%u70B9%u7EFF%u8335%u573A%u4E0A%u7684%u7CBE%u5F69%u77AC%u95F4",
"DNT": "1",
"Host": "tv.cctv.com",
"If-Modified-Since": "Sun, 16 Aug 2020 21:50:30 GMT",
"Sec-Fetch-Dest": "document",
"Sec-Fetch-Mode": "navigate",
"Sec-Fetch-Site": "none",
"Sec-Fetch-User": "?1",
"Upgrade-Insecure-Requests": "1",
"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/84.0.4147.125 Safari/537.36 Edg/84.0.522.59",
}


# In[4]:


r = requests.get(url = xwlb_url,headers = xwlb_headers)

r1 = requests.get(jsxsk)


# In[5]:


str1 = r.text

str1 = str1.encode('raw_unicode_escape').decode()

data = r1.text

data = data.encode('raw_unicode_escape').decode()


# In[6]:


soup = BeautifulSoup(str1,'html.parser')

soup1 = BeautifulSoup(data,'html.parser')


# In[7]:


def find_title(i):
    return soup.find_all(name = 'div', attrs = {'class':'title'})[i]

date = time.strftime("%Y-%m-%d", time.localtime())
date1 = datetime.datetime.now().strftime('%Y{y}%m{m}%d{d}').format(y='年', m='月', d='日')
 
data={}
data['标题']=str(find_title(0).string)
for i in range(1,11):
        data[str(i)]=str(find_title(i).string)


# In[8]:


raw = soup1.find_all(name = 'a', alt = True)


# In[9]:


doclocation='{} 新闻标题.docx'.format(date)


# In[10]:


Doc = Document()
Doc.styles['Normal'].font.name = u'宋体'
Doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
Doc.styles['Normal'].font.size = Pt(10.5)
Doc.styles['Normal'].font.color.rgb = RGBColor(0,0,0)

Doc.add_paragraph('标题:' + data['标题'])
for i in range(1,11):
        Doc.add_paragraph(str(i) + ':' + data[str(i)])
            


# In[11]:


Doc.add_heading(date)
for j in range(0,23):
    if raw[j].find('em').string==date1:
        Doc.add_paragraph(str(j+1)+':'+raw[j].find('span').string)
        
Doc.save(doclocation)

